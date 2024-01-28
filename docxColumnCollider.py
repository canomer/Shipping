from docx import Document
import os

def merge_third_and_fourth_columns(docx_file):
    # Belgeyi yükle
    doc = Document(docx_file)

    total_tables = len(doc.tables)
    if total_tables == 0:
        print("Bu belgede tablo bulunmuyor.")
        return

    for i, table in enumerate(doc.tables):
        for row in table.rows:
            # Her satır için 3. ve 4. sütunları kontrol et
            if len(row.cells) > 3:
                # 3. ve 4. sütunları birleştir
                a = row.cells[2]  # 3. sütun
                b = row.cells[3]  # 4. sütun
                a.merge(b)

        # İlerlemeyi yüzde olarak konsola yazdır
        progress = (i + 1) / total_tables * 100
        print(f"İşlem tamamlandı: %{progress:.2f}")

    # Değişiklikleri kaydet
    new_file = os.path.splitext(docx_file)[0] + "_modified.docx"
    doc.save(new_file)
    print(f"Değişiklikler kaydedildi: {new_file}")

# Scripti kullanmak için
# merge_third_and_fourth_columns("yolu/belgenizin_adı.docx")
