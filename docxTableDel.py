from docx import Document
import os

def clear_fourth_column(docx_file):
    # Belgeyi yükle
    doc = Document(docx_file)

    total_tables = len(doc.tables)
    if total_tables == 0:
        print("Bu belgede tablo bulunmuyor.")
        return

    for i, table in enumerate(doc.tables):
        # Her tablo için dördüncü sütunu kontrol et ve içeriğini temizle
        for row in table.rows:
            if len(row.cells) > 3:  # Eğer dördüncü sütun varsa
                row.cells[3].text = ''  # Dördüncü sütunun içeriğini temizle

        # İlerlemeyi yüzde olarak konsola yazdır
        progress = (i + 1) / total_tables * 100
        print(f"İşlem tamamlandı: %{progress:.2f}")

    # Değişiklikleri kaydet
    new_file = os.path.splitext(docx_file)[0] + "_modified.docx"
    doc.save(new_file)
    print(f"Değişiklikler kaydedildi: {new_file}")

# Scripti kullanmak için
# clear_fourth_column("yolu/belgenizin_adı.docx")
